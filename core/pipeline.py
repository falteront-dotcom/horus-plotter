"""
High-level pipeline: image → optimized G-code + preview.

Handles:
  - Auto-invert: if image is mostly dark, flip so content is drawn
  - Fit to work area: scale drawing to fit within plotter bounds
  - Center: offset so drawing is centered in work area
  - PDF/DOCX: convert first page to image before processing
"""

from __future__ import annotations

import os
from typing import TYPE_CHECKING

from PIL import Image, ImageStat

from .exceptions import EmptyDocumentError, ImageLoadError, UnsupportedFormatError
from .paths import Drawing, optimize
from .preprocessing import preprocess_image

if TYPE_CHECKING:
    from .gcode import PlotConfig


def _should_auto_invert(img: Image.Image) -> bool:
    """Return True if image is mostly dark and should be inverted."""
    stat = ImageStat.Stat(img)
    mean = stat.mean[0]
    return mean < 128


def _extract_docx_images(path: str) -> Image.Image | None:
    """Extract the first embedded image from a DOCX file.

    Returns grayscale PIL Image or None if no images found.
    """
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    import io

    doc = Document(path)

    # Collect all image blobs from relationships
    for rel in doc.part.rels.values():
        if rel.reltype == RT.IMAGE:
            try:
                blob = rel.target_part.blob
                img = Image.open(io.BytesIO(blob))
                return img.convert("L")
            except Exception:
                continue

    # Fallback: try inline shapes
    for shape in doc.inline_shapes:
        try:
            blip = shape._inline.graphic.graphicData
            for child in blip.iterchildren():
                if hasattr(child, 'blipFill'):
                    embed = child.blipFill.blip.get('embed')
                    if embed:
                        rel = doc.part.rels[embed]
                        blob = rel.target_part.blob
                        return Image.open(io.BytesIO(blob)).convert("L")
        except Exception:
            continue

    return None


def _render_docx_text(path: str) -> Image.Image:
    """Render DOCX text content as an image (fallback when no embedded images)."""
    from docx import Document
    from PIL import ImageDraw, ImageFont

    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    if not lines:
        # Empty document — return a blank white image
        return Image.new("L", (800, 1100), 255)

    # Render text onto image
    img = Image.new("L", (800, 1100), 255)
    draw = ImageDraw.Draw(img)
    y = 20
    try:
        font = ImageFont.truetype("arial.ttf", 18)
    except OSError:
        font = ImageFont.load_default()

    for line in lines[:50]:  # max 50 lines
        draw.text((20, y), line, fill=0, font=font)
        y += 24
        if y > 1080:
            break

    return img


def _load_as_image(path: str) -> Image.Image:
    """Load any supported file as a grayscale PIL Image.

    Supports: PNG, JPG, BMP, WebP, GIF, TIFF, PDF, DOCX.

    DOCX strategy:
      1. Extract embedded images (most reliable, no external deps)
      2. Try docx2pdf → PyMuPDF (requires MS Word)
      3. Render text as image (last resort)

    Args:
        path: Path to the input file.

    Returns:
        Grayscale PIL Image.

    Raises:
        ImageLoadError: If the file cannot be loaded.
        EmptyDocumentError: If the document contains no content.
        UnsupportedFormatError: If the file format is not supported.
    """
    ext = os.path.splitext(path)[1].lower()
    supported_image_exts = {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".gif"}

    if ext == ".pdf":
        return _load_pdf(path)

    if ext == ".docx":
        return _load_docx(path)

    if ext in (".tiff", ".tif"):
        return _load_tiff(path)

    if ext in supported_image_exts:
        return Image.open(path).convert("L")

    raise UnsupportedFormatError(path, ext)


def _load_pdf(path: str) -> Image.Image:
    """Load first page of a PDF as grayscale image."""
    try:
        import fitz  # pymupdf
    except ImportError as exc:
        raise ImageLoadError(path, "PyMuPDF (fitz) is not installed") from exc

    try:
        doc = fitz.open(path)
    except Exception as exc:
        raise ImageLoadError(path, f"Cannot open PDF: {exc}") from exc

    if len(doc) == 0:
        doc.close()
        raise EmptyDocumentError(path)

    try:
        page = doc[0]
        # Render at 300 DPI for high quality
        mat = fitz.Matrix(300 / 72, 300 / 72)
        pix = page.get_pixmap(matrix=mat)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        return img.convert("L")
    except Exception as exc:
        raise ImageLoadError(path, f"Failed to render PDF page: {exc}") from exc
    finally:
        doc.close()


def _load_docx(path: str) -> Image.Image:
    """Load a DOCX file as grayscale image."""
    # Strategy 1: extract embedded images
    img = _extract_docx_images(path)
    if img is not None:
        return img

    # Strategy 2: try docx2pdf (requires MS Word on Windows)
    try:
        from docx2pdf import convert as docx2pdf_convert  # type: ignore
        import tempfile
        tmp_pdf = os.path.join(tempfile.gettempdir(), "horus_docx.pdf")
        docx2pdf_convert(path, tmp_pdf)
        if os.path.exists(tmp_pdf):
            result = _load_pdf(tmp_pdf)
            try:
                os.remove(tmp_pdf)
            except OSError:
                pass
            return result
    except ImportError:
        pass  # docx2pdf not available
    except Exception as exc:
        # Log but continue to fallback
        import logging
        logging.getLogger(__name__).debug(
            "docx2pdf conversion failed: %s", exc
        )

    # Strategy 3: render text content as image
    return _render_docx_text(path)


def _load_tiff(path: str) -> Image.Image:
    """Load first frame of a TIFF as grayscale image."""
    try:
        img = Image.open(path)
        img.seek(0)
        return img.convert("L")
    except Exception as exc:
        raise ImageLoadError(path, f"Failed to load TIFF: {exc}") from exc


def image_to_gcode(image_path: str, cfg: PlotConfig) -> tuple[str, str]:
    """Full pipeline: image → G-code string + base64 preview PNG.

    Args:
        image_path: Path to the input image.
        cfg: Plot configuration.

    Returns:
        Tuple of (gcode_string, preview_base64).
    """
    from .gcode import drawing_to_gcode, drawing_to_preview
    from .converters import CONVERTERS, hatching

    img = _load_as_image(image_path)

    # Auto-invert: detect dark-background images
    if cfg.auto_invert and not cfg.invert:
        if _should_auto_invert(img):
            # Create a copy with invert=True
            import dataclasses
            cfg = dataclasses.replace(cfg, invert=True)

    # Preprocess: invert, brightness, contrast, blur
    img = preprocess_image(img, invert=cfg.invert,
                           brightness=cfg.brightness,
                           contrast=cfg.contrast,
                           blur=cfg.blur)

    # Calculate dimensions that fit within work area
    aspect = img.height / img.width
    max_w = cfg.work_area_x
    max_h = cfg.work_area_y

    # Start from requested width, but clamp if height exceeds work area
    draw_w = min(cfg.width_mm, max_w)
    draw_h = draw_w * aspect

    if draw_h > max_h:
        # Height exceeds work area — scale down to fit
        draw_h = max_h
        draw_w = draw_h / aspect

    # Build brightness grid
    grid_cols = max(2, round(draw_w / cfg.spacing_mm))
    grid_rows = max(2, round(draw_h / cfg.spacing_mm))
    img_small = img.resize((grid_cols, grid_rows), Image.Resampling.LANCZOS)
    pixels = img_small.load()

    # brightness grid: 0=black, 1=white
    grid = [
        [pixels[x, y] / 255.0 for x in range(grid_cols)]
        for y in range(grid_rows)
    ]

    # 1. Extract paths
    converter = CONVERTERS.get(cfg.style, hatching)

    # Styles that need width_mm and height_mm
    size_styles = ("spiral", "flow-field", "meandering", "concentric")
    if cfg.style in size_styles:
        drawing = converter(grid, cfg.spacing_mm, draw_w, draw_h,
                           threshold=cfg.threshold)
    elif cfg.style == "cross-hatching":
        drawing = converter(grid, cfg.spacing_mm, threshold=cfg.threshold)
    else:
        drawing = converter(grid, cfg.spacing_mm, threshold=cfg.threshold)

    # 2. Optimize paths
    drawing = optimize(drawing,
                       join_tolerance=cfg.join_tolerance,
                       simplify_epsilon=cfg.simplify_epsilon)

    # 3. Center the drawing in the work area
    #    Find actual bounding box (converters can produce negative coords)
    if drawing.paths:
        all_points = (
            pt for path in drawing.paths for pt in path.points
        )
        xs, ys = zip(*all_points)
        min_x, max_x = min(xs), max(xs)
        min_y, max_y = min(ys), max(ys)
        actual_w = max_x - min_x
        actual_h = max_y - min_y
        # Center the bounding box in the work area
        offset_x = (max_w - actual_w) / 2 - min_x
        offset_y = (max_h - actual_h) / 2 - min_y
    else:
        offset_x = 0.0
        offset_y = 0.0

    # 4. Convert to G-code (with centering offset)
    gcode = drawing_to_gcode(drawing, cfg, offset_x=offset_x, offset_y=offset_y)

    # 5. Render preview (A4-like sheet with paper background)
    preview = drawing_to_preview(drawing, work_area_x=max_w, work_area_y=max_h)

    return gcode, preview
